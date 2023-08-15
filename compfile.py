import time
import re
import docx as dx

def config():
    # 配置
    # 投标文件路径
    file1 = "C:/Users/Administrator/Desktop/招标文件/投标书v2.docx"
    file2 = 'C:/Users/Administrator/Desktop/招标文件/投标书 (副本)0719.docx'
    file3 = 'C:/Users/Administrator/Desktop/招标文件/投标书--欣诚信息技术有限公司.docx'

    # 招标文件路径
    doc_org2 = dx.Document('C:/Users/Administrator/Desktop/招标文件/招标文件2.docx')
    doc_org1 = dx.Document('C:/Users/Administrator/Desktop/招标文件/招标文件1.docx')

    # 计算的最小字句长度
    limitnum = 10

    # 分割词
    splitword = "。|:|：|,|，"

# 获取文档中的所有短句
def get_sentence(doc, splitword="", limitnum=10, needrun=False):
    for p in doc.paragraphs:
        for run in p.runs:
            text = run.text
            text = text.replace(" ", "").replace("\t", "")
            if len(splitword)>0:
                for t in re.split(splitword, text):
                    if len(t) >= limitnum:
                        if needrun:
                            yield [t, run]
                        else:
                            yield t
            else:
                if needrun:
                    yield [text, run]
                else:
                    yield text
    for table in doc.tables:
        for row in table.rows:
            try:
                row.cells[0].text
            except:
                continue
            for cell in row.cells:
                text = cell.text
                text = text.replace(" ", "").replace("\t", "")
                if len(splitword) > 0:
                    for t in re.split(splitword, text):
                        if len(t) >= limitnum:
                            if needrun:
                                yield [t, run]
                            else:
                                yield t
                else:
                    if needrun:
                        yield [text, run]
                    else:
                        yield text


def compare(doc_orgs, files, limitnum, splitword, colors, print):
    # 开始时间
    starttime = time.time()
    # 获取招标文件所有字句
    print("读取招标文件……")
    text_org = set()
    doc_org_all = list()
    for docfile in doc_orgs:
        doc = dx.Document(docfile)
        doc_org_all.append(doc)
        for t in get_sentence(doc, splitword, limitnum):
            text_org.add(t)
    # 获取招标文件中连续文本
    senten = ""
    for doc in doc_org_all:
        for text in get_sentence(doc, limitnum=1):
            senten = senten + text
    # 获取文档内容
    print("读取投标文件……")
    text_list = list()
    doc_all = list()
    for i in range(len(files)):
        print("第%s篇"%(i+1))
        text_set = set()
        doc = dx.Document(files[i])
        doc_all.append(doc)
        for t in get_sentence(doc, splitword, limitnum):
            text_set.add(t)
        text_list.append(text_set)

    readtime = time.time()
    print("文件读取完成,用时%s秒" % (readtime - starttime))
    # 寻找重复字句
    print("搜索重复字句")
    text_error_list = list()
    for n in range(len(text_list)):
        text_error = set()
        text_n = text_list[n]
        for text in text_n:
            if len(text) < limitnum:
                continue
            if text in text_org:
                continue
            for i in range(len(text_list)):
                if i!=n:
                    if text in text_list[i]:
                        text_error.add(text)
                        break
        text_error_ = set()
        for text in text_error:
            if text in senten:
                continue
            else:
                text_error_.add(text)
        text_error_list.append(text_error_)
    errortime = time.time()
    print("寻找重复字句完毕，用时%s秒" % (errortime - readtime))
    # 文档标记
    print("将重复段落标记成对应颜色…")
    for i in range(len(doc_all)):
        print('处理第%s个文档'%(i+1))
        doc = doc_all[i]
        for [text, run] in get_sentence(doc, splitword, limitnum, needrun=True):
            if text not in text_error_list[i]:
                continue
            if text in text_org:
                continue
            for j in range(len(doc_all)):
                if j!=i:
                    if text in text_error_list[j]:
                        run.font.color.rgb = dx.shared.RGBColor(colors[j][0], colors[j][1], colors[j][2])
                        break
        try:
            doc.save(files[i].replace(".docx", "_输出.docx"))
        except PermissionError:
            print("文档存储被拒绝，请确认是否已在office中关闭此文件。将在30s后重新尝试保存。")
            time.sleep(30)
            try:
                doc.save(files[i].replace(".docx", "_输出.docx"))
            except:
                print("保存失败！")
                print("【任务结束】")
                return
        with open(files[i].replace(".docx", "_重复子句.txt"), 'w') as f:
            for text in text_error_list[i]:
                try:
                    f.write(text + '\n')
                except:
                    pass

    endtimeall = time.time()
    print("docx文件染色完成，用时%s" % (endtimeall - errortime))
    print("【处理完成！】")



